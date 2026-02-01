"""DOCX document parser using python-docx."""

from pathlib import Path
import io

from src.nlp.ingestion.parser import ParsedDocument


class DocxParser:
    """Parser for DOCX documents."""

    def supports(self, filename: str) -> bool:
        """Check if this parser supports the file."""
        suffix = Path(filename).suffix.lower()
        return suffix in [".docx", ".doc"]

    async def parse(self, content: bytes, filename: str) -> ParsedDocument:
        """
        Parse DOCX document.
        
        Args:
            content: DOCX file content
            filename: Original filename
            
        Returns:
            ParsedDocument with extracted text and metadata
        """
        from docx import Document

        doc = Document(io.BytesIO(content))

        # Extract metadata
        core_props = doc.core_properties
        metadata = {
            "filename": filename,
            "title": core_props.title or "",
            "author": core_props.author or "",
            "subject": core_props.subject or "",
            "created": str(core_props.created) if core_props.created else "",
            "modified": str(core_props.modified) if core_props.modified else "",
        }

        # Extract text from paragraphs
        paragraphs = []
        sections = []
        current_section = {"heading": None, "content": []}

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            paragraphs.append(text)

            # Track sections based on heading styles
            if para.style.name.startswith("Heading"):
                if current_section["content"]:
                    sections.append(current_section)
                current_section = {
                    "heading": text,
                    "heading_level": self._get_heading_level(para.style.name),
                    "content": [],
                }
            else:
                current_section["content"].append(text)

        # Add last section
        if current_section["content"] or current_section["heading"]:
            sections.append(current_section)

        # Extract tables
        tables = self._extract_tables(doc)

        return ParsedDocument(
            text="\n\n".join(paragraphs),
            metadata=metadata,
            sections=sections if sections else None,
            tables=tables if tables else None,
        )

    def _get_heading_level(self, style_name: str) -> int:
        """Extract heading level from style name."""
        try:
            return int(style_name.replace("Heading", "").strip())
        except ValueError:
            return 1

    def _extract_tables(self, doc) -> list[dict]:
        """Extract tables from DOCX."""
        tables = []

        for i, table in enumerate(doc.tables):
            table_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    row_data.append(cell.text.strip())
                table_data.append(row_data)

            tables.append({
                "table_index": i,
                "data": table_data,
            })

        return tables
